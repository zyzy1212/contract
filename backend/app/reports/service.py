from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO

import fitz
from docx import Document

from app.contracts.service import EvidenceDTO, ReviewJobDetail


@dataclass(frozen=True)
class ReportDocument:
    text: str
    docx_bytes: bytes
    pdf_bytes: bytes


def _evidence_line(evidence: EvidenceDTO) -> str:
    snapshot = evidence.source_snapshot or {}
    title = str(snapshot.get("title") or "证据来源")
    article = str(snapshot.get("article_number") or "")
    page = str(snapshot.get("page_start") or snapshot.get("page") or "")
    version = str(snapshot.get("version") or "")
    source_url = str(snapshot.get("source_url") or "")
    line = f"《{title}》{article}"
    if page:
        line += f"，第 {page} 页"
    if version:
        line += f"，{version}"
    if source_url:
        line += f"，{source_url}"
    return line


def _report_lines(detail: ReviewJobDetail) -> list[str]:
    lines = [
        f"任务：{detail.id}",
        f"状态：{detail.status}",
        "",
    ]
    for finding in detail.findings:
        lines.extend(
            [
                f"风险：{finding.risk_level}",
                f"条款：{finding.clause_id}",
                f"问题：{finding.problem}",
                f"理由：{finding.reason}",
                f"建议：{finding.suggestion}",
                f"参考条款：{finding.proposed_clause}",
            ]
        )
        for evidence in finding.evidence:
            lines.append(f"证据：{_evidence_line(evidence)}")
        lines.append("")
    return lines


def _docx_bytes(lines: list[str]) -> bytes:
    document = Document()
    document.add_heading("合同审核报告", level=1)
    for line in lines:
        document.add_paragraph(line)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _pdf_bytes(lines: list[str]) -> bytes:
    document = fitz.open()
    page = document.new_page()
    y = 50
    for line in lines:
        if y > 780:
            page = document.new_page()
            y = 50
        page.insert_text((50, y), line, fontsize=10, fontname="china-s")
        y += 16
    data = document.tobytes()
    document.close()
    return data


def build_report(detail: ReviewJobDetail) -> ReportDocument:
    lines = _report_lines(detail)
    return ReportDocument(
        text="\n".join(lines),
        docx_bytes=_docx_bytes(lines),
        pdf_bytes=_pdf_bytes(lines),
    )
